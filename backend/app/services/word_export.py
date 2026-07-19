from __future__ import annotations

import io
import re
import uuid
from dataclasses import dataclass
from datetime import UTC, datetime
from urllib.parse import quote

from docx import Document as WordDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from sqlalchemy import select
from sqlalchemy.orm import Session, selectinload

from app.models.content import BulletPoint, Chapter, KnowledgePoint
from app.models.enums import ContentVersionKind
from app.models.project import GenerationJob, Project
from app.schemas.content import BulletPointRead, KnowledgePointRead
from app.services.content_organization import canonical_learning_material, is_displayable_source_chapter


VERSION_CONTRACT = (
    (ContentVersionKind.ORIGINAL, "原文版本"),
    (ContentVersionKind.RECITATION, "背诵版本"),
    (ContentVersionKind.KEYWORDS, "关键词版本"),
)
VERSION_LABELS = dict(VERSION_CONTRACT)
ALL_VERSION_ORDER = tuple(kind for kind, _ in VERSION_CONTRACT)
_UNSAFE_FILENAME = re.compile(r'[<>:"/\\|?*：／＼\x00-\x1f]+')
_ORDERED_ITEM = re.compile(r"^\s*(?:\d+[.、)）]|[（(][一二三四五六七八九十百\d]+[）)])\s*(.+)$")
_BULLET_ITEM = re.compile(r"^\s*[-–—•·]\s*(.+)$")
_SECRET = re.compile(r"(?i)(?:sk-[A-Za-z0-9_-]{8,}|(?:password|secret|token|api[_ -]?key)\s*[:=]\s*\S+)")


@dataclass(frozen=True)
class WordStyleTokens:
    page_width_cm: float = 21.0
    page_height_cm: float = 29.7
    margin_cm: float = 2.1
    header_footer_cm: float = 1.25
    latin_font: str = "Aptos"
    east_asia_font: str = "微软雅黑"
    body_size_pt: float = 10.5
    title_size_pt: float = 20.0
    version_size_pt: float = 16.0
    chapter_size_pt: float = 14.0
    knowledge_size_pt: float = 12.0
    internal_size_pt: float = 11.0
    meta_size_pt: float = 9.0
    body_line_spacing: float = 1.35
    body_after_pt: float = 6.0
    list_after_pt: float = 3.0
    ink: str = "25282B"
    muted: str = "6B7280"
    accent: str = "365F7D"


WORD_STYLE = WordStyleTokens()
STYLE_VERSION = "Revia Version Section"
STYLE_CHAPTER = "Revia Chapter"
STYLE_KNOWLEDGE = "Revia Knowledge Point"
STYLE_INTERNAL = "Revia Internal Heading"
STYLE_META = "Revia Metadata"


class WordExportNotFoundError(LookupError):
    pass


def safe_export_filename(project_name: str, version_label: str, exported_at: datetime) -> str:
    cleaned = _UNSAFE_FILENAME.sub("-", project_name).strip(" .-") or "Revia学习材料"
    cleaned = re.sub(r"\s+", " ", cleaned)[:80].rstrip(" .-")
    return f"{cleaned}-{version_label}-{exported_at:%Y-%m-%d}.docx"


def content_disposition(filename: str) -> str:
    ascii_fallback = re.sub(r"[^A-Za-z0-9._-]+", "-", filename).strip("-") or "revia-export.docx"
    return f"attachment; filename=\"{ascii_fallback}\"; filename*=UTF-8''{quote(filename)}"


class WordExportService:
    def __init__(self, db: Session) -> None:
        self._db = db

    def export(
        self,
        workspace_id: uuid.UUID,
        project_id: uuid.UUID,
        selected: ContentVersionKind | None,
        exported_at: datetime | None = None,
    ) -> tuple[io.BytesIO, str]:
        project = self._db.scalar(
            select(Project)
            .where(Project.id == project_id, Project.workspace_id == workspace_id)
            .options(
                selectinload(Project.chapters)
                .selectinload(Chapter.knowledge_points)
                .selectinload(KnowledgePoint.bullet_points)
                .selectinload(BulletPoint.versions),
                selectinload(Project.chapters)
                .selectinload(Chapter.knowledge_points)
                .selectinload(KnowledgePoint.bullet_points)
                .selectinload(BulletPoint.sources),
            )
        )
        if project is None:
            raise WordExportNotFoundError("项目不存在")
        now = exported_at or datetime.now(UTC)
        document = WordDocument()
        self.configure_document_styles(document)
        document.core_properties.title = project.name
        document.core_properties.subject = "Revia 学习材料"
        version_label = "全部版本" if selected is None else VERSION_LABELS[selected]
        self.add_document_title(document, project.name, now)
        chapters = canonical_learning_material(project.id, project.chapters).chapters
        export_kinds = ALL_VERSION_ORDER if selected is None else (selected,)
        for section_index, kind in enumerate(export_kinds):
            self.add_version_section(
                document,
                VERSION_LABELS[kind],
                page_break=section_index > 0,
            )
            for chapter in chapters:
                chapter_title = chapter.title
                if chapter_title:
                    self.add_chapter(document, chapter_title)
                for knowledge_point in sorted(chapter.knowledge_points, key=lambda item: item.position):
                    self.add_knowledge_point(document, knowledge_point, kind)

        failures = self._latest_failures(project.id)
        if failures:
            document.add_page_break()
            self.add_chapter(document, "未生成条目附录")
            for failure in failures:
                title = str(failure.get("syllabus_item") or "未命名条目")[:160]
                reason = _SECRET.sub("[敏感信息已隐藏]", str(failure.get("reason") or "未提供失败原因"))[:500]
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.add_run(title).bold = True
                paragraph.add_run(f"：{reason}")

        payload = io.BytesIO()
        document.save(payload)
        payload.seek(0)
        return payload, safe_export_filename(project.name, version_label, now)

    def _latest_failures(self, project_id: uuid.UUID) -> list[dict[str, object]]:
        job = self._db.scalar(
            select(GenerationJob)
            .where(GenerationJob.project_id == project_id)
            .order_by(GenerationJob.created_at.desc())
            .limit(1)
        )
        return list(job.item_failures or []) if job is not None else []

    @staticmethod
    def _stable_chapter_title(title: str | None) -> str | None:
        value = re.sub(r"\s+", " ", title or "").strip()
        return value if is_displayable_source_chapter(value) else None

    @classmethod
    def configure_document_styles(cls, document: WordDocument) -> None:
        section = document.sections[0]
        section.page_width = Cm(WORD_STYLE.page_width_cm)
        section.page_height = Cm(WORD_STYLE.page_height_cm)
        section.top_margin = Cm(WORD_STYLE.margin_cm)
        section.bottom_margin = Cm(WORD_STYLE.margin_cm)
        section.left_margin = Cm(WORD_STYLE.margin_cm)
        section.right_margin = Cm(WORD_STYLE.margin_cm)
        section.header_distance = Cm(WORD_STYLE.header_footer_cm)
        section.footer_distance = Cm(WORD_STYLE.header_footer_cm)

        cls._configure_paragraph_style(
            document, "Normal", WORD_STYLE.body_size_pt,
            color=WORD_STYLE.ink, after=WORD_STYLE.body_after_pt,
            line_spacing=WORD_STYLE.body_line_spacing,
        )
        cls._configure_paragraph_style(
            document, "Title", WORD_STYLE.title_size_pt,
            color=WORD_STYLE.ink, bold=True, before=18, after=8,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        cls._configure_paragraph_style(
            document, STYLE_VERSION, WORD_STYLE.version_size_pt,
            color=WORD_STYLE.accent, bold=True, before=10, after=14,
            keep_with_next=True,
        )
        cls._configure_paragraph_style(
            document, STYLE_CHAPTER, WORD_STYLE.chapter_size_pt,
            color=WORD_STYLE.ink, bold=True, before=18, after=8,
            keep_with_next=True,
        )
        cls._configure_paragraph_style(
            document, STYLE_KNOWLEDGE, WORD_STYLE.knowledge_size_pt,
            color=WORD_STYLE.ink, bold=True, before=12, after=6,
            keep_with_next=True,
        )
        cls._configure_paragraph_style(
            document, STYLE_INTERNAL, WORD_STYLE.internal_size_pt,
            color=WORD_STYLE.ink, bold=True, before=8, after=4,
            keep_with_next=True,
        )
        cls._configure_paragraph_style(
            document, STYLE_META, WORD_STYLE.meta_size_pt,
            color=WORD_STYLE.muted, after=4,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )
        for list_style in ("List Bullet", "List Number"):
            cls._configure_paragraph_style(
                document, list_style, WORD_STYLE.body_size_pt,
                color=WORD_STYLE.ink, after=WORD_STYLE.list_after_pt,
                line_spacing=WORD_STYLE.body_line_spacing,
            )
            paragraph_format = document.styles[list_style].paragraph_format
            paragraph_format.left_indent = Cm(0.8)
            paragraph_format.first_line_indent = Cm(-0.4)
        cls.add_page_number(section)

    @classmethod
    def _configure_paragraph_style(
        cls,
        document: WordDocument,
        name: str,
        size: float,
        *,
        color: str,
        bold: bool = False,
        before: float = 0,
        after: float = 0,
        line_spacing: float | None = None,
        alignment: WD_ALIGN_PARAGRAPH | None = None,
        keep_with_next: bool = False,
    ) -> None:
        style = document.styles[name] if name in document.styles else document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = WORD_STYLE.latin_font
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        fonts.set(qn("w:ascii"), WORD_STYLE.latin_font)
        fonts.set(qn("w:hAnsi"), WORD_STYLE.latin_font)
        fonts.set(qn("w:eastAsia"), WORD_STYLE.east_asia_font)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = keep_with_next
        style.paragraph_format.widow_control = True
        if line_spacing is not None:
            style.paragraph_format.line_spacing = line_spacing
        if alignment is not None:
            style.paragraph_format.alignment = alignment

    @staticmethod
    def add_page_number(section) -> None:
        paragraph = section.footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.font.name = WORD_STYLE.latin_font
        run.font.size = Pt(WORD_STYLE.meta_size_pt)
        run.font.color.rgb = RGBColor.from_string(WORD_STYLE.muted)
        fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
        fonts.set(qn("w:ascii"), WORD_STYLE.latin_font)
        fonts.set(qn("w:hAnsi"), WORD_STYLE.latin_font)
        fonts.set(qn("w:eastAsia"), WORD_STYLE.east_asia_font)
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = " PAGE "
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        value = OxmlElement("w:t")
        value.text = "1"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.extend((begin, instruction, separate, value, end))

    @staticmethod
    def add_document_title(document: WordDocument, project_name: str, exported_at: datetime) -> None:
        document.add_paragraph(project_name, style="Title")
        document.add_paragraph(
            f"导出日期 {exported_at.astimezone().strftime('%Y-%m-%d')}",
            style=STYLE_META,
        )

    @staticmethod
    def add_version_section(document: WordDocument, label: str, *, page_break: bool) -> None:
        if page_break:
            document.add_page_break()
        document.add_paragraph(label, style=STYLE_VERSION)

    @staticmethod
    def add_chapter(document: WordDocument, title: str) -> None:
        document.add_paragraph(title, style=STYLE_CHAPTER)

    def add_knowledge_point(
        self,
        document: WordDocument,
        knowledge_point: KnowledgePointRead,
        kind: ContentVersionKind,
    ) -> None:
        document.add_paragraph(knowledge_point.title, style=STYLE_KNOWLEDGE)
        for bullet in sorted(knowledge_point.bullet_points, key=lambda item: item.position):
            self._add_bullet_version(document, bullet, kind)

    def _add_bullet_version(self, document: WordDocument, bullet: BulletPointRead, kind: ContentVersionKind) -> None:
        version = next((item for item in bullet.versions if item.kind == kind), None)
        if version is None:
            return
        document.add_paragraph(version.title, style=STYLE_INTERNAL)
        self._add_content(document, version.content)

    @staticmethod
    def _add_content(document: WordDocument, content: str) -> None:
        for raw_block in re.split(r"\n\s*\n", content.strip()):
            lines = [line.strip() for line in raw_block.splitlines() if line.strip()]
            if not lines:
                continue
            for line in lines:
                ordered = _ORDERED_ITEM.match(line)
                bullet = _BULLET_ITEM.match(line)
                if ordered:
                    document.add_paragraph(ordered.group(1), style="List Number")
                elif bullet:
                    document.add_paragraph(bullet.group(1), style="List Bullet")
                else:
                    document.add_paragraph(line)
