import unittest
import uuid
import zipfile
from datetime import UTC, datetime

from docx import Document as WordDocument
from docx.oxml.ns import qn
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool

import app.models  # noqa: F401
from app.db.base import Base
from app.models.content import BulletPoint, BulletPointSource, Chapter, ContentVersion, KnowledgePoint
from app.models.enums import ContentVersionKind, GenerationStatus
from app.models.project import GenerationJob, Project
from app.models.workspace import Workspace
from app.services.word_export import (
    STYLE_CHAPTER,
    STYLE_INTERNAL,
    STYLE_KNOWLEDGE,
    STYLE_META,
    STYLE_VERSION,
    WORD_STYLE,
    WordExportNotFoundError,
    WordExportService,
)


class WordExportTests(unittest.TestCase):
    def setUp(self) -> None:
        self.engine = create_engine(
            "sqlite+pysqlite:///:memory:",
            connect_args={"check_same_thread": False},
            poolclass=StaticPool,
        )
        self.Session = sessionmaker(bind=self.engine, expire_on_commit=False)
        Base.metadata.create_all(self.engine)
        self.workspace_id = uuid.uuid4()
        self.project_id = uuid.uuid4()
        self.first_source_id = uuid.uuid4()
        self.second_source_id = uuid.uuid4()
        with self.Session() as db:
            db.add(Workspace(id=self.workspace_id))
            project = Project(id=self.project_id, workspace_id=self.workspace_id, name="财政/政策：复习")
            chapter = Chapter(title="资料章节：财政政策", position=0)
            knowledge = KnowledgePoint(title="财政政策工具", position=0)
            bullet = BulletPoint(position=0)
            bullet.versions = [
                ContentVersion(kind=ContentVersionKind.ORIGINAL, title="政策作用机制", content="原文版本正文。\n\n1. 政府支出\n2. 税收"),
                ContentVersion(kind=ContentVersionKind.RECITATION, title="政策作用机制", content="背诵版本考试表达。"),
                ContentVersion(kind=ContentVersionKind.KEYWORDS, title="政策作用机制", content="- 支出\n- 税收"),
            ]
            bullet.sources = [BulletPointSource(
                text_chunk_id=self.first_source_id,
                page_start=12,
                page_end=13,
            )]
            knowledge.bullet_points.append(bullet)
            chapter.knowledge_points.append(knowledge)
            project.chapters.append(chapter)
            second_chapter = Chapter(title="资料章节：货币政策", position=1)
            second_knowledge = KnowledgePoint(title="货币政策工具", position=0)
            second_bullet = BulletPoint(position=0)
            second_bullet.versions = [
                ContentVersion(kind=ContentVersionKind.ORIGINAL, title="公开市场操作", content="原文第二章正文。"),
                ContentVersion(kind=ContentVersionKind.RECITATION, title="公开市场操作", content="背诵第二章正文。"),
                ContentVersion(kind=ContentVersionKind.KEYWORDS, title="公开市场操作", content="公开市场、利率、货币供给"),
            ]
            second_bullet.sources = [BulletPointSource(
                text_chunk_id=self.second_source_id,
                page_start=90,
                page_end=90,
            )]
            second_knowledge.bullet_points.append(second_bullet)
            second_chapter.knowledge_points.append(second_knowledge)
            project.chapters.append(second_chapter)
            db.add(project)
            db.add(GenerationJob(
                project_id=self.project_id,
                status=GenerationStatus.PARTIAL_FAILED,
                provider="mock",
                item_failures=[{
                    "syllabus_item": "自动稳定器",
                    "reason": "token=secret-value 未找到足够资料依据",
                }],
            ))
            db.commit()

    def tearDown(self) -> None:
        self.engine.dispose()

    @staticmethod
    def _text(payload) -> tuple[WordDocument, str]:
        payload.seek(0)
        document = WordDocument(payload)
        return document, "\n".join(paragraph.text for paragraph in document.paragraphs)

    def test_current_version_exports_one_version_without_source_pages(self) -> None:
        with self.Session() as db:
            payload, filename = WordExportService(db).export(
                self.workspace_id,
                self.project_id,
                ContentVersionKind.ORIGINAL,
                datetime(2026, 7, 19, tzinfo=UTC),
            )
        self.assertTrue(zipfile.is_zipfile(payload))
        document, text = self._text(payload)
        self.assertIn("财政/政策：复习", text)
        self.assertIn("资料章节：财政政策", text)
        self.assertIn("财政政策工具", text)
        self.assertIn("政策作用机制", text)
        self.assertIn("原文版本正文", text)
        self.assertIn("原文第二章正文", text)
        self.assertNotIn("背诵版本考试表达", text)
        self.assertNotIn("关键词版本", text)
        self.assertIn("原文版本", text)
        self.assertNotIn("来源：", text)
        self.assertNotIn("第 12", text)
        self.assertIn("未生成条目附录", text)
        self.assertIn("自动稳定器", text)
        self.assertNotIn("secret-value", text)
        self.assertTrue(any(paragraph.style.name == "List Number" for paragraph in document.paragraphs))
        self.assertNotIn("/", filename)
        self.assertNotIn("：", filename)
        self.assertTrue(filename.endswith("-原文版本-2026-07-19.docx"))

    def test_all_versions_are_three_complete_sections_with_page_breaks(self) -> None:
        with self.Session() as db:
            payload, filename = WordExportService(db).export(self.workspace_id, self.project_id, None)
        with zipfile.ZipFile(payload) as archive:
            self.assertIn("[Content_Types].xml", archive.namelist())
            self.assertIn("word/document.xml", archive.namelist())
        document, text = self._text(payload)
        for label in ("原文版本", "背诵版本", "关键词版本"):
            self.assertIn(label, text)
            self.assertEqual(sum(paragraph.text == label for paragraph in document.paragraphs), 1)
        self.assertLess(text.index("原文版本"), text.index("背诵版本"))
        self.assertLess(text.index("背诵版本"), text.index("关键词版本"))
        self.assertLess(text.index("原文第二章正文"), text.index("背诵版本"))
        self.assertLess(text.index("背诵第二章正文"), text.index("关键词版本"))
        self.assertGreater(text.index("公开市场、利率、货币供给"), text.index("关键词版本"))
        self.assertIn("背诵版本考试表达", text)
        self.assertIn("全部版本", filename)
        self.assertEqual(text.count("财政/政策：复习"), 1)
        self.assertNotIn("来源：", text)
        self.assertEqual(len(document._element.xpath('.//w:br[@w:type="page"]')), 3)

    def test_styles_are_centralized_a4_and_chinese_font_is_explicit(self) -> None:
        with self.Session() as db:
            payload, _ = WordExportService(db).export(
                self.workspace_id,
                self.project_id,
                ContentVersionKind.ORIGINAL,
            )
        document, _ = self._text(payload)
        section = document.sections[0]
        self.assertAlmostEqual(section.page_width.cm, WORD_STYLE.page_width_cm, places=1)
        self.assertAlmostEqual(section.page_height.cm, WORD_STYLE.page_height_cm, places=1)
        for margin in (section.top_margin, section.bottom_margin, section.left_margin, section.right_margin):
            self.assertAlmostEqual(margin.cm, WORD_STYLE.margin_cm, places=1)

        expected_sizes = {
            "Title": WORD_STYLE.title_size_pt,
            STYLE_VERSION: WORD_STYLE.version_size_pt,
            STYLE_CHAPTER: WORD_STYLE.chapter_size_pt,
            STYLE_KNOWLEDGE: WORD_STYLE.knowledge_size_pt,
            STYLE_INTERNAL: WORD_STYLE.internal_size_pt,
            "Normal": WORD_STYLE.body_size_pt,
            STYLE_META: WORD_STYLE.meta_size_pt,
            "List Number": WORD_STYLE.body_size_pt,
            "List Bullet": WORD_STYLE.body_size_pt,
        }
        for name, expected_size in expected_sizes.items():
            style = document.styles[name]
            self.assertAlmostEqual(style.font.size.pt, expected_size)
            self.assertEqual(style._element.rPr.rFonts.get(qn("w:eastAsia")), WORD_STYLE.east_asia_font)
            self.assertEqual(style._element.rPr.rFonts.get(qn("w:ascii")), WORD_STYLE.latin_font)
        for name in (STYLE_VERSION, STYLE_CHAPTER, STYLE_KNOWLEDGE, STYLE_INTERNAL):
            self.assertTrue(document.styles[name].paragraph_format.keep_with_next)
        self.assertIn("PAGE", section.footer._element.xml)

    def test_docx_xml_contains_no_source_or_evidence_page_labels(self) -> None:
        with self.Session() as db:
            payload, _ = WordExportService(db).export(self.workspace_id, self.project_id, None)
        with zipfile.ZipFile(payload) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")
        for forbidden in ("来源：", "来源页码", "引用页码", "证据页码", "source references", "evidence pages"):
            self.assertNotIn(forbidden, xml)

    def test_unreliable_placeholder_chapter_labels_are_omitted(self) -> None:
        with self.Session() as db:
            project = db.get(Project, self.project_id)
            for position, title, body in (
                (2, "未分章", "未分章下的正文仍需导出。"),
                (3, "未归类内容", "未归类标题下的正文仍需导出。"),
                (4, "人力 · 第 16–114 页", "页码回退标题下的正文仍需导出。"),
            ):
                chapter = Chapter(title=title, position=position)
                knowledge = KnowledgePoint(title=f"占位标题测试 {position}", position=0)
                bullet = BulletPoint(position=0)
                bullet.versions = [ContentVersion(
                    kind=ContentVersionKind.ORIGINAL,
                    title="测试内容",
                    content=body,
                )]
                knowledge.bullet_points.append(bullet)
                chapter.knowledge_points.append(knowledge)
                project.chapters.append(chapter)
            db.commit()
            payload, _ = WordExportService(db).export(
                self.workspace_id,
                self.project_id,
                ContentVersionKind.ORIGINAL,
            )
        _, text = self._text(payload)
        self.assertIn("未分章下的正文仍需导出。", text)
        self.assertIn("未归类标题下的正文仍需导出。", text)
        self.assertIn("页码回退标题下的正文仍需导出。", text)
        self.assertNotIn("未分章\n", text)
        self.assertNotIn("未归类内容", text)
        self.assertNotIn("第 16–114 页", text)

    def test_cross_level_duplicate_is_removed_from_word_without_losing_complete_point(self) -> None:
        with self.Session() as db:
            project = db.get(Project, self.project_id)
            chapter = project.chapters[0]
            broad = KnowledgePoint(title="人力资源基础概念", position=1)
            complete = KnowledgePoint(title="人力资源的特征", position=2)

            def add_bullet(owner, title, content, position):
                item = BulletPoint(position=position)
                item.versions = [
                    ContentVersion(kind=ContentVersionKind.ORIGINAL, title=title, content=content),
                    ContentVersion(kind=ContentVersionKind.RECITATION, title=title, content=f"{content}背诵"),
                    ContentVersion(kind=ContentVersionKind.KEYWORDS, title=title, content=f"{title}、定义、要点"),
                ]
                owner.bullet_points.append(item)

            add_bullet(broad, "人力资源的特征", "只覆盖能动性、可再生性。", 0)
            add_bullet(broad, "人力资源的含义", "宽泛知识点的其他内容。", 1)
            for index, title in enumerate(("能动性", "可再生性", "增值性", "时效性", "社会性")):
                add_bullet(complete, title, f"{title}完整正文。", index)
            chapter.knowledge_points.extend((broad, complete))
            db.commit()
            payload, _ = WordExportService(db).export(
                self.workspace_id, self.project_id, ContentVersionKind.ORIGINAL,
            )
            self.assertEqual(len(broad.bullet_points), 2, "read-time canonicalization must not mutate ORM data")

        _, text = self._text(payload)
        self.assertEqual(text.count("人力资源的特征"), 1)
        self.assertNotIn("只覆盖能动性、可再生性。", text)
        for title in ("能动性", "可再生性", "增值性", "时效性", "社会性"):
            self.assertIn(f"{title}完整正文。", text)
        self.assertIn("宽泛知识点的其他内容。", text)

    def test_duplicate_knowledge_and_bullet_title_is_hidden_without_losing_content(self) -> None:
        with self.Session() as db:
            project = db.get(Project, self.project_id)
            knowledge = project.chapters[0].knowledge_points[0]
            for version in knowledge.bullet_points[0].versions:
                version.title = "第 1 节：财政政策工具"
            db.commit()
            payload, _ = WordExportService(db).export(
                self.workspace_id,
                self.project_id,
                ContentVersionKind.ORIGINAL,
            )

        document, text = self._text(payload)
        self.assertIn("原文版本正文", text)
        self.assertEqual(sum(paragraph.text == "财政政策工具" for paragraph in document.paragraphs), 1)
        self.assertNotIn("第 1 节：财政政策工具", text)

    def test_export_is_workspace_scoped(self) -> None:
        with self.Session() as db:
            with self.assertRaises(WordExportNotFoundError):
                WordExportService(db).export(uuid.uuid4(), self.project_id, None)


if __name__ == "__main__":
    unittest.main()
