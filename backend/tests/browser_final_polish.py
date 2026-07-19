import io
import json
import tempfile
import zipfile
from pathlib import Path
from urllib.parse import quote

from docx import Document
from playwright.sync_api import Route, sync_playwright


OWNER_PROJECT_ID = "10000000-0000-0000-0000-000000000001"
DEMO_PROJECT_ID = "20000000-0000-0000-0000-000000000002"


def project_payload(project_id: str, name: str):
    return {
        "id": project_id,
        "name": name,
        "description": "浏览器封版验证",
        "status": "completed",
        "created_at": "2026-07-19T00:00:00Z",
        "updated_at": "2026-07-19T00:00:00Z",
    }


def material_payload(project_id: str):
    chapters = []
    for chapter_index, chapter_title in enumerate(("资料第一章 市场结构", "资料第二章 政策工具", "资料第三章 综合应用")):
        points = []
        for point_index in range(3):
            bullet_id = f"{chapter_index + 3}{point_index + 3}000000-0000-0000-0000-000000000001"
            title = f"内部小标题 {chapter_index + 1}-{point_index + 1}"
            detailed = "\n\n".join([
                f"原文版本正文用于长内容阅读与搜索定位。章节 {chapter_index + 1}，知识点 {point_index + 1}。",
                "1. 第一项完整说明\n2. 第二项完整说明\n3. 第三项完整说明",
                "补充段落。" * 24,
            ])
            versions = [
                {"id": bullet_id[:-1] + "2", "kind": "original", "title": title, "content": detailed},
                {"id": bullet_id[:-1] + "3", "kind": "recitation", "title": title, "content": "背诵版本正文适合考试表达。\n\n- 定义\n- 机制\n- 结论"},
                {"id": bullet_id[:-1] + "4", "kind": "keywords", "title": title, "content": "关键词版本、记忆线索、中文检索"},
            ]
            points.append({
                "id": bullet_id[:-1] + "5",
                "title": f"知识点 {chapter_index + 1}-{point_index + 1} 中文定位",
                "position": point_index,
                "bullet_points": [{"id": bullet_id, "position": 0, "versions": versions, "sources": []}],
            })
        chapters.append({
            "id": f"{chapter_index + 6}0000000-0000-0000-0000-000000000001",
            "title": chapter_title,
            "position": chapter_index,
            "knowledge_points": points,
        })
    return {"project_id": project_id, "chapters": chapters}


def docx_bytes() -> bytes:
    document = Document()
    document.add_heading("独立演示项目", 0)
    document.add_heading("资料第一章 市场结构", 1)
    document.add_paragraph("背诵版本正文适合考试表达。")
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def install_api(context):
    state = {"role": None}
    exported = docx_bytes()

    def handler(route: Route):
        request = route.request
        path = request.url.split("/api/backend", 1)[-1].split("?", 1)[0]
        if path == "/auth/mode":
            route.fulfill(json={"public_access_enabled": False, "demo_access_enabled": True})
        elif path == "/auth/session":
            route.fulfill(status=401, json={"detail": "需要工作区授权"})
        elif path == "/auth/access":
            code = json.loads(request.post_data or "{}").get("access_code")
            state["role"] = "demo" if code == "demo-browser-code" else "owner"
            workspace_id = "00000000-0000-0000-0000-000000000202" if state["role"] == "demo" else "00000000-0000-0000-0000-000000000101"
            route.fulfill(json={"workspace_id": workspace_id, "role": state["role"]})
        elif path == "/auth/logout":
            route.fulfill(status=204, body="")
        elif path == "/projects/active-document":
            route.fulfill(body="null", headers={"content-type": "application/json"})
        elif path.endswith("/generation-jobs/latest"):
            route.fulfill(body="null", headers={"content-type": "application/json"})
        elif path.endswith("/learning-material"):
            project_id = path.split("/")[2]
            route.fulfill(json=material_payload(project_id))
        elif path.endswith("/exports/word"):
            filename = "独立演示项目-背诵版本-2026-07-19.docx"
            route.fulfill(
                body=exported,
                headers={
                    "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "content-disposition": f"attachment; filename*=UTF-8''{quote(filename)}",
                },
            )
        elif path == "/projects":
            if state["role"] == "demo":
                route.fulfill(json=[project_payload(DEMO_PROJECT_ID, "独立演示项目")])
            else:
                route.fulfill(json=[project_payload(OWNER_PROJECT_ID, "固定站长项目")])
        elif path.startswith("/projects/"):
            project_id = path.split("/")[2]
            name = "独立演示项目" if project_id == DEMO_PROJECT_ID else "固定站长项目"
            route.fulfill(json=project_payload(project_id, name))
        else:
            route.fulfill(status=404, json={"detail": f"unhandled browser fixture: {path}"})

    context.route("**/api/backend/**", handler)


def login(context, code: str, console_errors: list[str]):
    install_api(context)
    page = context.new_page()
    page.on("console", lambda message: console_errors.append(message.text) if message.type == "error" and "Failed to load resource" not in message.text else None)
    page.on("pageerror", lambda error: console_errors.append(str(error)))
    page.goto("http://127.0.0.1:3000")
    page.wait_for_load_state("networkidle")
    page.locator("#app-access-code").fill(code)
    page.get_by_role("button", name="进入 Workspace").click()
    page.locator("button.project-row").first.wait_for()
    return page


with sync_playwright() as playwright:
    browser = playwright.chromium.launch(headless=True)
    console_errors: list[str] = []

    owner_context = browser.new_context(viewport={"width": 1440, "height": 900})
    owner = login(owner_context, "owner-browser-code", console_errors)
    assert owner.get_by_text("固定站长项目", exact=True).count() >= 1

    second_context = browser.new_context(viewport={"width": 1280, "height": 800})
    second_owner = login(second_context, "owner-browser-code", console_errors)
    assert second_owner.get_by_text("固定站长项目", exact=True).count() >= 1
    second_context.close()

    demo_context = browser.new_context(viewport={"width": 1440, "height": 900}, accept_downloads=True)
    demo = login(demo_context, "demo-browser-code", console_errors)
    assert demo.get_by_text("演示模式 · 修改不会保存", exact=True).is_visible()
    assert demo.get_by_role("button", name="演示模式只读").is_disabled()
    demo.get_by_text("独立演示项目", exact=True).last.click()
    demo.locator(".reading-document").wait_for()
    assert demo.get_by_role("button", name="演示模式只读").is_disabled()

    search = demo.get_by_label("搜索当前项目")
    search.fill("资料第二章")
    demo.get_by_text("1 个结果", exact=True).wait_for()
    demo.locator(".search-results-panel li button").first.click()
    assert demo.locator(".is-search-target").count() >= 1
    demo.get_by_role("button", name="下一个结果").click()
    demo.get_by_role("button", name="上一个结果").click()
    demo.get_by_role("button", name="清空").click()
    assert demo.locator("mark").count() == 0

    demo.get_by_role("tab", name="背诵版本").click()
    search.fill("背诵版本正文")
    demo.get_by_text("9 个结果", exact=True).wait_for()
    demo.get_by_role("button", name="清空").click()

    reading = demo.locator(".reading-scroll")
    last_knowledge_id = demo.locator(".knowledge-section").last.get_attribute("id")
    reading.evaluate("(element, id) => { element.scrollTop = document.getElementById(id).offsetTop - 20 }", last_knowledge_id)
    demo.wait_for_timeout(500)
    assert demo.locator(".outline-nav button.is-active").count() >= 1

    first_bullet = demo.locator(".bullet-point").first
    first_bullet.click(button="right")
    demo.get_by_role("menuitem", name="编辑").hover()
    demo.get_by_role("menuitem", name="整体编辑").click()
    drawer = demo.locator(".keyword-drawer")
    assert drawer.is_visible()
    triggers = drawer.locator(".accordion-trigger")
    for index in range(3):
        if triggers.nth(index).get_attribute("aria-expanded") == "true":
            triggers.nth(index).click()
    assert drawer.locator('.accordion-trigger[aria-expanded="true"]').count() == 0
    triggers.nth(1).click()
    assert drawer.locator('.accordion-trigger[aria-expanded="true"]').count() == 1
    triggers.nth(0).click()
    triggers.nth(2).click()
    assert drawer.locator('.accordion-trigger[aria-expanded="true"]').count() == 3

    middle_editor = drawer.get_by_label("背诵版本内容")
    middle_editor.fill("\n\n".join([f"长编辑内容 {index}" for index in range(120)]))
    drawer_body = drawer.locator(".drawer-body")
    middle_editor.evaluate("element => { element.scrollTop = element.scrollHeight }")
    before = drawer_body.evaluate("element => element.scrollTop")
    middle_editor.hover()
    demo.mouse.wheel(0, 1400)
    demo.wait_for_timeout(200)
    after = drawer_body.evaluate("element => element.scrollTop")
    assert after > before
    drawer.get_by_role("button", name="临时应用全部").click()
    assert not drawer.is_visible()

    demo.get_by_role("button", name="导出").click()
    with demo.expect_download() as download_info:
        demo.get_by_role("button", name="导出当前版本").click()
    download = download_info.value
    output = Path(tempfile.gettempdir()) / download.suggested_filename
    download.save_as(output)
    assert output.suffix == ".docx"
    assert zipfile.is_zipfile(output)

    demo.set_viewport_size({"width": 700, "height": 760})
    first_bullet = demo.locator(".bullet-point").first
    first_bullet.click(button="right")
    demo.get_by_role("menuitem", name="编辑").hover()
    demo.get_by_role("menuitem", name="整体编辑").click()
    box = demo.locator(".keyword-drawer").bounding_box()
    assert box and box["width"] <= 700
    demo.screenshot(path="storage/browser-final-polish.png", full_page=True)

    assert not console_errors, console_errors
    demo_context.close()
    owner_context.close()
    browser.close()

print("BROWSER_FINAL_POLISH=passed")
